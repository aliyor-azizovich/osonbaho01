from PyQt5.QtWidgets import (
    QWidget, QCheckBox, QPushButton, QLabel, QDialog, QMessageBox, QTableWidgetItem,
    QLineEdit, QTextBrowser, QFileDialog
)
from PyQt5 import uic
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QFont

import os
import re
import sys
import json
import traceback

from num2words import num2words
from jinja2 import Environment, DebugUndefined, Template

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from bs4 import BeautifulSoup

from ui.payment_dialog import PaymaentDialog
from logic.license_checker import is_license_valid
from ui.method__dialog import MethodRejectionDialog

from ui.cost_method_dialogs.building_choose import BuildingChooseDialog
from ui.cost_method_dialogs.deviations_and_wear_dialog import DeviationsAndWearDialog

from logic.data_entry import DataEntryForm
from logic.paths import get_ui_path


class AgreementWidget(QWidget):
    """
    Упрощённая версия:
    - Единственный метод согласования: Среднее арифметическое
    - Никаких процентов/спинбоксов и комбобоксов
    - Итог: среднее двух активных подходов или одно значение, если активен только один
    """
    FIXED_METHOD = "Среднее арифметическое"

    def __init__(self, parent=None, main_window=None, valuation_window=None):
        super().__init__(parent)
        self.main_window = main_window
        self.valuation_window = valuation_window
        uic.loadUi(get_ui_path("agreement_widget.ui"), self)

        self.data_service = DataEntryForm()

        # UI элементы (минимально необходимые)
        self.label_agreement_method = self.findChild(QLabel, "label_agreement_method")  # может отсутствовать в старом ui
        if self.label_agreement_method:
            self.label_agreement_method.setText(f"Метод согласования: {self.FIXED_METHOD}")

        self.checkBox_cost = self.findChild(QCheckBox, 'checkBox_cost')
        self.checkBox_comparative = self.findChild(QCheckBox, 'checkBox_comparative')

        self.label_cost_weighted_average = self.findChild(QLabel, 'label_cost_weighted_average')
        self.label_comparative_weighted_average = self.findChild(QLabel, 'label_comparative_weighted_average')

        self.label_final_cost = self.findChild(QLabel, 'label_final_cost')
        self.label_building_land = self.findChild(QLabel, 'label_building_land')

        self.pushButton_final_save = self.findChild(QPushButton, 'pushButton_final_save')
        if self.pushButton_final_save:
            self.pushButton_final_save.clicked.connect(self.final_save)

        self.pushButton_upload_report = self.findChild(QPushButton, 'pushButton_upload_report')
        if self.pushButton_upload_report:
            self.pushButton_upload_report.clicked.connect(self.on_upload_report_clicked)

        # Стартовые состояния
        if self.checkBox_cost:
            self.checkBox_cost.setChecked(False)
            self.checkBox_cost.stateChanged.connect(self.update_agreement_ui)
        if self.checkBox_comparative:
            self.checkBox_comparative.setChecked(False)
            self.checkBox_comparative.stateChanged.connect(self.update_agreement_ui)

        # Данные
        self.cost_value = 0.0
        self.comparative_value = 0.0

        # Вспомогательные атрибуты
        self.building_cost = 0.0
        self.land_cost = 0.0

        self.update_agreement_ui()

    # ===== Вспомогательное =====

    def format_sum(self, value):
        try:
            return f"{round(float(value)):,}".replace(",", " ")
        except Exception:
            return str(value)

    # ===== Загрузка исходных стоимостей из общего JSON =====

    def load_costs_from_json(self, full_data):
        """
        Загружает суммы для Затратного и Сравнительного подходов и обновляет подписи.
        """
        try:
            # Затратный подход: сумма литеров + земля
            liters = full_data.get("liters", [])
            building_cost = sum(liter.get("final_cost", 0) for liter in liters)

            land_text = full_data.get("land_valuation", {}).get("land_total_cost", "")
            land_cost = 0.0
            if land_text:
                land_cost = float(
                    land_text.split(":")[-1]
                    .replace("сум", "")
                    .replace(" ", "")
                    .replace(",", "")
                )

            self.building_cost = building_cost
            self.land_cost = land_cost
            self.cost_value = building_cost + land_cost

            if self.checkBox_cost:
                self.checkBox_cost.setText(
                    f"Затратный подход: {self.format_sum(self.cost_value)} сум"
                )
            if self.label_building_land:
                self.label_building_land.setText(
                    f"Стоимость улучшений: {self.format_sum(building_cost)} сум\n"
                    f"Права на землю: {self.format_sum(land_cost)} сум"
                )

            # Сравнительный подход
            comp_text = full_data.get("comparative", {}).get("label_comparative_final_cost", "")
            comp_number = "".join(c for c in comp_text if c.isdigit() or c in ",.")
            self.comparative_value = float(comp_number.replace(" ", "").replace(",", "")) if comp_number else 0.0

            if self.checkBox_comparative:
                self.checkBox_comparative.setText(
                    f"Сравнительный подход: {self.format_sum(self.comparative_value)} сум"
                )

            # Обновить UI после загрузки
            self.update_agreement_ui()

        except Exception as e:
            print("[ERROR] Ошибка загрузки стоимостей из JSON:", e)

    # ===== Основная логика согласования (только среднее арифметическое) =====

    def update_agreement_ui(self):
        """
        Показываем итог:
        - Если активен 1 подход: показываем его сумму, скрываем подсчёты по долям
        - Если активны оба: итог = среднее, а под каждым показываем половину его значения
        """
        use_cost = self.checkBox_cost.isChecked() if self.checkBox_cost else False
        use_comparative = self.checkBox_comparative.isChecked() if self.checkBox_comparative else False

        # Сначала прячем вспомогательные лейблы
        if self.label_cost_weighted_average:
            self.label_cost_weighted_average.setVisible(False)
        if self.label_comparative_weighted_average:
            self.label_comparative_weighted_average.setVisible(False)

        active_count = sum([use_cost, use_comparative])

        if active_count == 0:
            if self.label_final_cost:
                self.label_final_cost.setText("Итоговая стоимость: н/д")
            return

        if active_count == 1:
            # Один подход
            if use_cost and self.cost_value:
                if self.label_final_cost:
                    self.label_final_cost.setText(
                        f"Итоговая стоимость: {self.format_sum(self.cost_value)} сум"
                    )
            elif use_comparative and self.comparative_value:
                if self.label_final_cost:
                    self.label_final_cost.setText(
                        f"Итоговая стоимость: {self.format_sum(self.comparative_value)} сум"
                    )
            return

        # Два активных подхода -> простое среднее
        total = 0.0
        cnt = 0
        if use_cost and self.cost_value is not None:
            total += float(self.cost_value)
            cnt += 1
        if use_comparative and self.comparative_value is not None:
            total += float(self.comparative_value)
            cnt += 1

        avg = total / cnt if cnt else 0.0
        if self.label_final_cost:
            self.label_final_cost.setText(
                f"Итоговая стоимость: {self.format_sum(avg)} сум"
            )

        # Показать «половинки» под каждым подходом
        if self.label_cost_weighted_average and use_cost:
            self.label_cost_weighted_average.setVisible(True)
            self.label_cost_weighted_average.setText(
                f"{self.format_sum(self.cost_value / 2)} сум"
            )
        if self.label_comparative_weighted_average and use_comparative:
            self.label_comparative_weighted_average.setVisible(True)
            self.label_comparative_weighted_average.setText(
                f"{self.format_sum(self.comparative_value / 2)} сум"
            )

    # ===== Подготовка данных для отчёта =====

    def collect_agreement_data(self):
        # финальный текст и число
        final_cost_text = self.label_final_cost.text() if self.label_final_cost else "Итоговая стоимость: н/д"
        clean_text = final_cost_text.replace('\xa0', ' ').replace('\u202f', ' ')
        match = re.search(r'([\d\s]+)\s*сум', clean_text.lower())
        edited_cost = match.group(1).strip() if match else ""

        try:
            words = num2words(int(edited_cost.replace(" ", "")), lang='ru').capitalize() + " сум"
        except Exception:
            words = ""

        use_cost = self.checkBox_cost.isChecked() if self.checkBox_cost else False
        use_comparative = self.checkBox_comparative.isChecked() if self.checkBox_comparative else False
        active_count = sum([use_cost, use_comparative])

        # Текстовое описание метода
        if active_count == 1:
            agreement_method_summary = (
                "Оценка рыночной стоимости объекта произведена на основе одного подхода, "
                "который признан наиболее достоверным и обоснованным с учётом специфики объекта "
                "и доступности данных. Согласование результатов не производилось."
            )
        else:
            agreement_method_summary = (
                "Согласование результатов оценки выполнено методом простого среднего, при котором каждый из выбранных "
                "подходов имеет равную значимость. Итоговая стоимость определена как среднее арифметическое значений."
            )

        # Стоимости
        building_cost_raw = getattr(self, "building_cost", 0.0)
        land_cost_raw = getattr(self, "land_cost", 0.0)
        total_cost_value_raw = getattr(self, "cost_value", 0.0)
        comparative_value_raw = getattr(self, "comparative_value", 0.0)

        # Проценты/взвешенные значения для отчёта
        if active_count == 2:
            cost_percent = 50
            comparative_percent = 50
            weighted_cost = total_cost_value_raw / 2
            weighted_comparative = comparative_value_raw / 2
        elif use_cost:
            cost_percent = 100
            comparative_percent = 0
            weighted_cost = total_cost_value_raw
            weighted_comparative = 0
        else:
            cost_percent = 0
            comparative_percent = 100
            weighted_cost = 0
            weighted_comparative = comparative_value_raw

        return {
            "method": self.FIXED_METHOD,
            "use_cost": use_cost,
            "use_comparative": use_comparative,
            "cost_percent": cost_percent,
            "comparative_percent": comparative_percent,
            "final_cost": final_cost_text,
            "edited_final_cost": edited_cost,
            "amount_in_words": words,
            "building_cost": round(building_cost_raw),
            "land_cost": round(land_cost_raw),
            "total_cost_value": round(total_cost_value_raw),
            "comparative_final_cost_value": round(comparative_value_raw),
            "agreement_method_summary": agreement_method_summary,
            "weighted_cost": round(weighted_cost),
            "weighted_comparative": round(weighted_comparative),
            "weights": {
                "cost": cost_percent,
                "comparative": comparative_percent
            }
        }

    def load_agreement_data(self, data):
        """
        Восстанавливаем состояние вкладки (без комбобокса и спинбоксов).
        """
        self._loading = True
        try:
            if self.label_agreement_method:
                self.label_agreement_method.setText(f"Метод согласования: {self.FIXED_METHOD}")

            # чекбоксы
            if self.checkBox_cost:
                self.checkBox_cost.setChecked(data.get("use_cost", True))
            if self.checkBox_comparative:
                self.checkBox_comparative.setChecked(data.get("use_comparative", True))

            # значения
            self.cost_value = float(data.get("total_cost_value", 0) or 0)
            self.comparative_value = float(data.get("comparative_final_cost_value", 0) or 0)

            if self.checkBox_cost:
                self.checkBox_cost.setText(f"Затратный подход: {self.format_sum(self.cost_value)} сум")
            if self.checkBox_comparative:
                self.checkBox_comparative.setText(f"Сравнительный подход: {self.format_sum(self.comparative_value)} сум")

            # детали по затратному
            self.building_cost = float(data.get("building_cost", 0) or 0)
            self.land_cost = float(data.get("land_cost", 0) or 0)
            if self.label_building_land:
                self.label_building_land.setText(
                    f"Стоимость улучшений: {self.format_sum(self.building_cost)} сум\n"
                    f"Права на землю: {self.format_sum(self.land_cost)} сум"
                )

            # итог
            if self.label_final_cost:
                self.label_final_cost.setText(data.get("final_cost", "Итоговая стоимость: н/д"))

            # обновить вычисления по текущим чекбоксам
            self.update_agreement_ui()

        except Exception as e:
            print(f"[ERROR] Не удалось загрузить данные вкладки Согласование: {e}")
        finally:
            self._loading = False

    # ===== Действия кнопок =====

    def final_save(self):
        self.valuation_window.save_report()
        self.valuation_window.main_window.update_last_valuation_cost_from_agreement()

    # ===== Ниже — функции формирования отчёта (оставлены как были) =====

    def insert_kadastr_table_into_word(self, docx_path, html_table, marker="{{TABLE_KADASTR}}"):
        soup = BeautifulSoup(html_table, "html.parser")
        table_html = soup.find("table")
        if not table_html:
            return
        rows = table_html.find_all("tr")
        if not rows:
            return

        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                p = paragraph._element
                p_parent = p.getparent()

                table = doc.add_table(rows=0, cols=2)
                tbl_element = table._element
                p.addnext(tbl_element)

                try:
                    table.style = 'Table Grid'
                except:
                    pass

                for tr in rows:
                    cols = tr.find_all("td")
                    if len(cols) != 2:
                        continue
                    row_cells = table.add_row().cells
                    for i in range(2):
                        text = cols[i].get_text(strip=True)
                        row_cells[i].text = text
                        for par in row_cells[i].paragraphs:
                            par.alignment = WD_ALIGN_PARAGRAPH.CENTER

                p_parent.remove(p)
                break
        doc.save(docx_path)

    def load_market_analysis(self, oblast_name, research_folder="research"):
        import os
        docx_filename = f"{oblast_name}.docx"
        base_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
        docx_path = os.path.join(base_dir, research_folder, docx_filename)
        if not os.path.exists(docx_path):
            return f"⚠️ Анализ по области «{oblast_name}» не найден."
        doc = Document(docx_path)
        full_text = [para.text for para in doc.paragraphs]
        return "\n\n".join(full_text)

    def insert_comparative_table(self, docx_path, comparative, marker="[ comparative_table ]"):
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                parent.remove(paragraph._element)

                vertical_headers = comparative.get("vertical_headers", [])
                horizontal_headers = comparative.get("horizontal_headers", [])
                table_data = comparative.get("table_data", [])
                final_value = comparative.get("label_comparative_final_cost", "—")

                if not vertical_headers or not horizontal_headers or not table_data:
                    break

                cols = len(horizontal_headers) + 1
                table = doc.add_table(rows=0, cols=cols)
                table.style = "Table Grid"
                header_row = table.add_row().cells
                header_row[0].text = "Параметры"
                for i, analog in enumerate(horizontal_headers):
                    header_row[i + 1].text = analog["text"]

                for row_index, row_label in enumerate(vertical_headers[:-1]):
                    row_cells = table.add_row().cells
                    row_cells[0].text = row_label
                    for col_index in range(len(horizontal_headers)):
                        value = table_data[row_index][col_index]
                        row_cells[col_index + 1].text = value

                final_row = table.add_row().cells
                merged_cell = final_row[0].merge(final_row[-1])
                paragraph = merged_cell.paragraphs[0]
                run = paragraph.add_run(final_value)
                run.bold = True
                run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                parent.insert(index, table._element)
                break
        doc.save(docx_path)

    def improved_insert_land_table(self, docx_path, land_valuation, marker="[ LAND_TABLE ]"):
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                parent.remove(paragraph._element)

                analogs = land_valuation["horizontal_headers"]
                vertical_headers = land_valuation["vertical_headers"]
                table_data = land_valuation["table_data"]

                table = doc.add_table(rows=0, cols=len(analogs) + 1)
                table.style = "Table Grid"

                header_row = table.add_row().cells
                header_row[0].text = "Параметры"
                for i, analog in enumerate(analogs):
                    header_row[i + 1].text = analog["text"]

                for row_index, row_label in enumerate(vertical_headers[:-1]):
                    row_cells = table.add_row().cells
                    row_cells[0].text = row_label
                    for col_index in range(len(analogs)):
                        value = table_data[row_index][col_index]
                        row_cells[col_index + 1].text = value

                avg_row = table.add_row().cells
                merged = avg_row[0].merge(avg_row[-1])
                p = merged.paragraphs[0]
                run = p.add_run(
                    "Средняя стоимость за 1 сотку: " + land_valuation.get("cost_per_sotka_soum", "—")
                )
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                parent.insert(index, table._element)
                break
        doc.save(docx_path)

    def insert_agreement_table(self, docx_path, agreement_data, marker='[ agreement_table ]'):
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                parent.remove(paragraph._element)

                use_cost = agreement_data.get("use_cost", False)
                use_comparative = agreement_data.get("use_comparative", False)

                if use_cost and use_comparative:
                    table = doc.add_table(rows=2, cols=7)
                    table.style = "Table Grid"

                    headers = [
                        "Затратный подход", "", "",
                        "Сравнительный подход", "", "",
                        "Согласованная стоимость"
                    ]
                    subheaders = [
                        "стоимость улучшений + стоимость права на земельный участок", "Удельный вес", "Взвешенное значение",
                        "Стоимость объекта", "Удельный вес", "Взвешенное значение",
                        ""
                    ]
                    for i, text in enumerate(headers):
                        cell = table.cell(0, i); cell.text = text
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for i, text in enumerate(subheaders):
                        cell = table.cell(1, i); cell.text = text
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    def fmt(v):
                        try:
                            return f"{round(float(str(v).replace(' ', ''))):,}".replace(",", " ")
                        except:
                            return str(v)

                    row = table.add_row().cells
                    row[0].text = fmt(agreement_data.get("total_cost_value", 0))
                    row[1].text = str(agreement_data.get("cost_percent", ""))
                    row[2].text = fmt(agreement_data.get("weighted_cost", 0))
                    row[3].text = fmt(agreement_data.get("comparative_final_cost_value", 0))
                    row[4].text = str(agreement_data.get("comparative_percent", ""))
                    row[5].text = fmt(agreement_data.get("weighted_comparative", 0))
                    row[6].text = fmt(agreement_data.get("edited_final_cost", 0))

                else:
                    table = doc.add_table(rows=3, cols=3)
                    table.style = "Table Grid"
                    headers = ["Объект оценки", "", "Стоимость согласования"]
                    subheaders = ["", "Стоимость", ""]
                    for i, text in enumerate(headers):
                        table.cell(0, i).text = text
                        table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for i, text in enumerate(subheaders):
                        table.cell(1, i).text = text
                        table.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    method = "Затратный подход" if use_cost else "Сравнительный подход"
                    cost = agreement_data.get("edited_final_cost", "")
                    row = table.rows[2].cells
                    row[0].text = "Жилой дом"
                    row[1].text = cost
                    row[2].text = cost

                parent.insert(index, table._element)
                break
        doc.save(docx_path)

    def insert_koeff_table(self, docx_path, koefs_data, marker="[ koeff_table ]"):
        doc = Document(docx_path)
        koefs_table = koefs_data.get("koefs_table", [])
        if not koefs_table:
            return
        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                parent.remove(paragraph._element)

                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                header = table.rows[0].cells
                header[0].text = "Дата"
                header[1].text = "Коэффициент"

                for row in koefs_table:
                    if len(row) != 2:
                        continue
                    date, value = row
                    cells = table.add_row().cells
                    cells[0].text = date
                    cells[1].text = value

                parent.insert(index, table._element)
                break
        doc.save(docx_path)

    # ===== Выгрузка отчёта =====

    def on_upload_report_clicked(self):
        self.valuation_window.save_report()

        agreement = self.collect_agreement_data()
        use_cost = agreement.get("use_cost", False)
        use_comparative = agreement.get("use_comparative", False)

        if not use_cost or not use_comparative:
            dialog = MethodRejectionDialog(self, valuation_window=self.valuation_window)
            if dialog.exec_() == QDialog.Accepted:
                self.selected_template = dialog.selected_template
            else:
                QMessageBox.warning(self, "Отмена", "Выгрузка отменена.")
                return

        valuation_window = self.valuation_window
        report_number = valuation_window.report_number_input.text().strip()
        report_path = valuation_window.main_window.report_manager.get_report_path(report_number)

        if not os.path.exists(report_path):
            QMessageBox.warning(self, "Ошибка", f"Файл отчёта не найден:\n{report_path}")
            return

        with open(report_path, "r", encoding="utf-8") as f:
            report_data = json.load(f)

        agreement = report_data.get("agreement", {})
        use_cost = agreement.get("use_cost", True)
        use_comparative = agreement.get("use_comparative", True)

        template_name = getattr(self, "selected_template", "result.docx")
        template_path = os.path.join("reports", template_name)
        if not os.path.exists(template_path):
            QMessageBox.critical(self, "Ошибка", f"Файл шаблона {template_name} не найден.")
            return

        if not is_license_valid():
            QMessageBox.warning(
                self,
                "Подписка истекла",
                "Срок действия подписки истёк или не подтверждён.\nПожалуйста, оплатите подписку."
            )
            dialog = PaymaentDialog(self)
            result = dialog.exec_()
            if not is_license_valid():
                return

        def generate_engineering_description(comms: dict, heating: str) -> str:
            mapping = {
                "газификация": "Газоснабжение",
                "электроосвещение": "Электроосвещение",
                "водоснабжение": "Водоснабжение",
                "канализация": "Канализация",
                "телефонная_линия": "Телефонная линия",
                "электрический_водонагреватель": "Электрический водонагреватель",
                "горячее_водоснабжение": "Горячее водоснабжение"
            }
            available = [name for key, name in mapping.items() if comms.get(key)]
            count = len(available)
            if count <= 2:
                level = "низкая"
            elif count <= 4:
                level = "удовлетворительная"
            else:
                level = "хорошая"

            lines = ["Инженерная инфраструктура:"]
            if available:
                lines.append("Объект обеспечен следующими инженерными коммуникациями:")
                lines.extend(f"– {item}" for item in available)
            else:
                lines.append("Объект не обеспечен инженерными коммуникациями.")
            lines.append("")
            lines.append(f"Отопление: {heating or 'не указано'}")
            lines.append(f"Степень обеспеченности инженерной инфраструктурой — {level}.")
            return "\n".join(lines)

        def insert_liter_tables(docx_path, liters):
            doc = Document(docx_path)

            def remove_last_paragraph_if_empty(doc_):
                if doc_.paragraphs and not doc_.paragraphs[-1].text.strip():
                    p_ = doc_.paragraphs[-1]._element
                    p_.getparent().remove(p_)

            def insert_full_liter(doc_, parent_, index_, liter):
                def add_row(table_, key, value):
                    row = table_.add_row().cells
                    row[0].text = key
                    row[1].text = str(value)

                def format2(val):
                    return f"{float(val):.2f}" if val not in ("", None) else ""

                table = doc_.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                table.autofit = True

                filters = liter.get("filters", {})
                m = liter.get("measurements", {})

                add_row(table, "Тип строения", liter["building_type"])
                if filters.get("Этажность"): add_row(table, "Этажность", filters["Этажность"])
                if filters.get("Отделка"): add_row(table, "Отделка", filters["Отделка"])
                if filters.get("Толщина стен"): add_row(table, "Толщина стен", filters["Толщина стен"])
                if filters.get("Материал стен"): add_row(table, "Материал стен", filters["Материал стен"])
                if filters.get("Примыкание"): add_row(table, "Примыкание", filters["Примыкание"])
                if filters.get("Фундамент"): add_row(table, "Фундамент", filters["Фундамент"])
                if filters.get("Кровля"): add_row(table, "Кровля", filters["Кровля"])
                if filters.get("Перекрытие"): add_row(table, "Перекрытие", filters["Перекрытие"])
                if filters.get("Тип покрытия"): add_row(table, "Тип покрытия", filters["Тип покрытия"])
                if filters.get("Высота стен"): add_row(table, "Высота стен", filters["Высота стен"])
                if m.get("square"): add_row(table, "Площадь", m["square"])
                if m.get("length"): add_row(table, "Протяжённость", m["length"])
                if m.get("height"): add_row(table, "Высота", m["height"])
                if m.get("volume"): add_row(table, "Объём", m["volume"])
                if m.get("ukup_price_label"): add_row(table, "Стоимость по УКУП", m["ukup_price_label"])

                remove_last_paragraph_if_empty(doc_)
                parent_.insert(index_, table._element)
                index_ += 1

                html = liter.get("analog_description_html", "")
                if html:
                    soup = BeautifulSoup(html, "html.parser")
                    desc_title = doc_.add_paragraph("")
                    parent_.insert(index_, desc_title._element)
                    index_ += 1
                    paragraphs = soup.get_text(separator="\n").split("\n")
                    for text in paragraphs:
                        if text.strip():
                            para = doc_.add_paragraph(text.strip())
                            parent_.insert(index_, para._element)
                            index_ += 1

                analog_table = doc_.add_table(rows=1, cols=3)
                analog_table.style = "Table Grid"
                analog_table.autofit = True

                header = analog_table.rows[0].cells
                header[0].text = "Наименование"
                header[1].text = "Значение"
                header[2].text = "Комментарий"

                def add_analog_row(name, value, comment):
                    row = analog_table.add_row().cells
                    row[0].text = name
                    row[1].text = str(value)
                    row[2].text = comment

                facade_price = liter.get("facade_corrected_price", 0)
                if facade_price:
                    add_analog_row("Корректировка на фасад", format2(facade_price), liter.get("facade_type", ""))

                improvements = liter.get("improvement_details", [])
                if improvements:
                    for imp in improvements:
                        name = imp.get("name", "Неизвестное улучшение")
                        percent = imp.get("correction_percent", 0.0)
                        absolute = imp.get("correction_value", 0.0)
                        add_analog_row(name, f"{absolute:,.2f} сум", f"{percent * 100:+.2f}%")

                deviations = liter.get("deviation_details", [])
                if deviations:
                    for dev in deviations:
                        if dev.get("selected", False):
                            name = dev.get("name", "Отклонение")
                            value = dev.get("value", 0)
                            add_analog_row(name, f"{value:+,.2f} сум", "Отклонение")

                height_correction = liter.get("height_corrected_price", 0)
                if height_correction:
                    add_analog_row("Корректировка по высоте", format2(height_correction), "сум")

                add_analog_row("Откорректированная стоимость УКУП", format2(liter["corrected_price"]), "сум")
                add_analog_row(liter["reg_coeff_type"], liter["reg_coeff"], "")
                add_analog_row(liter["stat_koeff_label"], f"{liter['stat_coeff']:.3f}", "")
                add_analog_row("Прибыль предпринимателя", f"{liter['developer_percent']}", "%")
                add_analog_row("Коэффициент сейсмичности", f"{liter['sesmos']}", "")
                add_analog_row("Восттановительная стоимость", f"{self.format_sum(liter['replacement_cost'])}", "Сум")

                remove_last_paragraph_if_empty(doc_)
                parent_.insert(index_, analog_table._element)
                index_ += 1

                para3 = doc_.add_paragraph("РАСЧЁТ ФИЗИЧЕСКОГО ИЗНОСА", style="Heading 3")
                parent_.insert(index_, para3._element)
                index_ += 1

                se_table = doc_.add_table(rows=1, cols=4)
                se_table.style = "Table Grid"
                se_table.autofit = True
                se_header = se_table.rows[0].cells
                se_header[0].text = "Конструктивные элементы"
                se_header[1].text = "Доля %"
                se_header[2].text = "Поправка"
                se_header[3].text = "Износ %"

                for se in liter["structural_elements"]:
                    row = se_table.add_row().cells
                    row[0].text = se["Конструкции"]
                    row[1].text = str(se["Доля %"])
                    row[2].text = str(se.get("Поправка к удельным весам %") or se.get("Поправка к\nудельным весам %") or "—")
                    row[3].text = str(se.get("Физический износ %") or se.get("Физический\nизнос %") or "—")

                remove_last_paragraph_if_empty(doc_)
                parent_.insert(index_, se_table._element)
                index_ += 1

                summary_table = doc_.add_table(rows=4, cols=2)
                summary_table.style = "Table Grid"
                summary_table.autofit = True
                summary_table.rows[0].cells[0].text = "Коэф. недостроенности"
                summary_table.rows[0].cells[1].text = str(liter.get("inconsistency", ""))
                summary_table.rows[1].cells[0].text = "Общий процент износа"
                summary_table.rows[1].cells[1].text = f"{liter['wear_percent']:,.2f} %"
                summary_table.rows[2].cells[0].text = "Физический износ"
                summary_table.rows[2].cells[1].text = f"{liter['wear_price']:,.2f} сум"
                summary_table.rows[3].cells[0].text = "Оценочная стоимость"
                cell = summary_table.rows[3].cells[1]
                cell.text = ""
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(f"{liter['final_cost']:,.2f} сум")
                run.bold = True

                remove_last_paragraph_if_empty(doc_)
                parent_.insert(index_, summary_table._element)
                index_ += 1

                empty = doc_.add_paragraph()
                parent_.insert(index_, empty._element)
                index_ += 1
                return index_

            for paragraph in doc.paragraphs:
                if "[[LITER_TABLES_PLACEHOLDER]]" in paragraph.text:
                    parent = paragraph._element.getparent()
                    index = parent.index(paragraph._element)
                    parent.remove(paragraph._element)
                    for liter in liters:
                        para = doc.add_paragraph(f"Литер №{liter['number']}", style='Heading 2')
                        parent.insert(index, para._element)
                        index += 1
                        index = insert_full_liter(doc, parent, index, liter)
                    break
            doc.save(docx_path)

        try:
            valuation_window = self.valuation_window
            report_number = valuation_window.report_number_input.text().strip()
            reg_number = valuation_window.lineEdit_reg_number.text().strip() or f"Report_{report_number}"

            report_path = valuation_window.main_window.report_manager.get_report_path(report_number)
            if not os.path.exists(report_path):
                QMessageBox.warning(self, "Ошибка", f"Файл отчёта не найден:\n{report_path}")
                return

            with open(report_path, "r", encoding="utf-8") as f:
                report_data = json.load(f)

            template_path = os.path.join("reports", getattr(self, "selected_template", "result.docx"))
            if not os.path.exists(template_path):
                QMessageBox.critical(self, "Ошибка", "Файл шаблона result.docx не найден.")
                return

            save_dir = valuation_window.main_window.save_directory
            if not save_dir:
                QMessageBox.critical(self, "Ошибка", "Путь сохранения не задан.")
                return

            report_folder = os.path.join(save_dir, reg_number)
            os.makedirs(report_folder, exist_ok=True)
            output_path = os.path.join(report_folder, f"Отчёт №{report_number}.docx")

            liters = report_data.get("liters", [])
            lines = []
            for liter in liters:
                number = liter.get("number", "")
                building_type = liter.get("building_type", "")
                measurements = liter.get("measurements", {})
                square = measurements.get("square", "")
                length = measurements.get("length", "")
                if building_type.lower().startswith("огражд") or building_type.lower() == "забор":
                    value = length or "—"; unit = "м"; label = "протяжённость"
                else:
                    value = square or "—"; unit = "м²"; label = "площадь"
                lines.append(f"№{number}. {building_type}, {label}: {value} {unit}")
            liters_block_text = "\n".join(lines)

            def sanitize_html_table(raw_html):
                cleaned = raw_html.replace('\n', ' ').replace('\r', ' ')
                cleaned = ' '.join(cleaned.split())
                cleaned = cleaned.replace("‘", "'").replace("’", "'").replace("`", "'")
                cleaned = cleaned.replace("“", '"').replace("”", '"')
                cleaned = cleaned.replace("*", "")
                return cleaned

            raw_table_html = sanitize_html_table(report_data.get("kadastr_table_html", ""))
            rendered_table_html = Template("{{ val | safe }}").render(val=raw_table_html)

            oblast_name = report_data.get("administrative", {}).get("oblast", "")
            regional_market_analysis = self.load_market_analysis(oblast_name)

            communications = report_data.get("communications", {})
            heating = report_data.get("heating", "")
            engineering_description = generate_engineering_description(communications, heating)

            rayon = report_data.get("administrative", {}).get("rayon", "").lower()
            density = "высокая" if "город" in rayon else "средняя"

            agreement = report_data.get("agreement", {})
            use_cost = agreement.get("use_cost", False)
            use_comparative = agreement.get("use_comparative", False)

            context = {
                "liters_block": liters_block_text,
                "liters": report_data.get("liters", []),
                "report_number": report_number,
                "reg_number": reg_number,
                "contract_date": report_data.get("contract_date", ""),
                "inspection_date": report_data.get("inspection_date", ""),
                "exchange_rate": report_data.get("exchange_rate", ""),
                "address": report_data.get("address", ""),
                "owner_name": report_data.get("owner_name", ""),
                "valuation_purpose": report_data.get("valuation_purpose", ""),
                "price_type": report_data.get("price_type", ""),
                "buyer_type": report_data.get("buyer_type", ""),
                "buyer_name": report_data.get("buyer_name", ""),
                "buyer_passport_series": report_data.get("buyer_passport_series", ""),
                "buyer_passport_number": report_data.get("buyer_passport_number", ""),
                "buyer_address": report_data.get("buyer_address", ""),
                "land_area": report_data.get("land_area", ""),
                "total_area": report_data.get("total_area", ""),
                "useful_area": report_data.get("useful_area", ""),
                "living_area": report_data.get("living_area", ""),
                "cadastral_number": report_data.get("cadastral_number", ""),
                "profit": report_data.get("profit", ""),
                "agreement": {
                    "method": self.FIXED_METHOD,
                    "cost_percent": report_data.get("agreement", {}).get("cost_percent", ""),
                    "comparative_percent": report_data.get("agreement", {}).get("comparative_percent", ""),
                    "edited_final_cost": report_data.get("agreement", {}).get("edited_final_cost", ""),
                    "amount_in_words": report_data.get("agreement", {}).get("amount_in_words", ""),
                    "building_cost": (
                        self.format_sum(agreement.get("building_cost") or 0) if use_cost else "Затратный подход не использовался"
                    ),
                    "land_cost": (
                        self.format_sum(agreement.get("land_cost", "")) if use_cost else "Затратный подход не использовался"
                    ),
                    "total_cost_value": (
                        self.format_sum(agreement.get("total_cost_value", "")) if use_cost else "Затратный подход не использовался"
                    ),
                    "comparative_final_cost_value": (
                        self.format_sum(agreement.get("comparative_final_cost_value", "")) if use_comparative else "Сравнительный подход не использовался"
                    ),
                },
                'analogs_count': report_data.get("land_valuation", {}).get('analogs_count', ""),
                "administrative": {
                    "oblast": report_data.get("administrative", {}).get("oblast", ""),
                    "rayon": report_data.get("administrative", {}).get("rayon", "")
                },
                "regional_market_analysis": regional_market_analysis,
                "engineering_description": engineering_description,
                "density": density,
                "lineEdit_CBUF": report_data.get("lineEdit_CBUF", ""),
                "profit": report_data.get("profit"),
            }

            context["TABLE_KADASTR"] = "{{ TABLE_KADASTR }}"
            context["agreement_method_summary"] = report_data["agreement"].get("agreement_method_summary", "")

            env = Environment(undefined=DebugUndefined)
            doc = DocxTemplate(template_path)
            doc.env = env
            doc.render(context)
            doc.save(output_path)

            insert_liter_tables(output_path, report_data["liters"])
            self.improved_insert_land_table(output_path, report_data["land_valuation"], marker="[ LAND_TABLE ]")
            self.insert_koeff_table(output_path, report_data["koefs"], marker="[ koeff_table ]")
            self.insert_comparative_table(output_path, report_data['comparative'], marker="[ comparative_table ]")
            self.insert_kadastr_table_into_word(output_path, raw_table_html, marker="{{ TABLE_KADASTR }}")
            self.insert_agreement_table(output_path, report_data["agreement"], marker='[ agreement_table ]')

            QMessageBox.information(self, "Успех", f"Отчёт успешно создан:\n{output_path}")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сгенерировать отчёт:\n{str(e)}")
